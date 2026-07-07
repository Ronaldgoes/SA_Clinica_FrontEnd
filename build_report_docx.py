from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
ASSETS_DIR = BASE_DIR / "word-assets"
OUTPUT_DOCX = BASE_DIR / "RELATORIO_FEATURES.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def style_run(run, *, bold=False, size=None, color=None, italic=False):
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Arial"


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    style_run(run, size=10.5)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    style_run(run, bold=True, size=15, color="0F4C5C")


def add_body_paragraph(doc, text, *, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    style_run(run, size=10.5, italic=italic)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    style_run(run, size=9, color="64748B", italic=True)


def add_image(doc, image_path, width=Inches(6.6)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(image_path), width=width)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(10.5)

for style_name, size, color in [
    ("Title", 24, "0F172A"),
    ("Heading 1", 16, "0F4C5C"),
    ("Heading 2", 13, "134E4A"),
    ("Heading 3", 11.5, "1E293B"),
]:
    style = styles[style_name]
    style.font.name = "Arial"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(18)
title.paragraph_format.space_after = Pt(6)
run = title.add_run("Relatorio de Melhorias - Clinica")
style_run(run, bold=True, size=24, color="0F172A")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(10)
run = subtitle.add_run("Entrega com apoio de IA, documentacao das features e prints da aplicacao")
style_run(run, size=11, color="475569")

meta = doc.add_table(rows=2, cols=2)
meta.style = "Table Grid"
meta.autofit = False
meta.columns[0].width = Inches(1.7)
meta.columns[1].width = Inches(4.8)
meta.rows[0].cells[0].text = "Projeto"
meta.rows[0].cells[1].text = "Clinica"
meta.rows[1].cells[0].text = "Data"
meta.rows[1].cells[1].text = "30/06/2026"
for row in meta.rows:
    for idx, cell in enumerate(row.cells):
        set_cell_margins(cell)
        if idx == 0:
            set_cell_shading(cell, "D9EEF3")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    style_run(r, bold=True, size=10)
        else:
            for p in cell.paragraphs:
                for r in p.runs:
                    style_run(r, size=10)

doc.add_paragraph()

add_section_heading(doc, "1. Objetivo")
add_body_paragraph(
    doc,
    "Implementar melhorias no sistema de clinica sem quebrar as funcionalidades existentes, mantendo o padrao visual do projeto e registrando os principais pontos da implementacao.",
)

add_section_heading(doc, "2. Features entregues")
add_bullet(doc, "Sistema de busca avancada por nome, convenio, alergias, telefone e identificador.")
add_bullet(doc, "Ordenacao de consultas e exames por data, com alternancia entre mais recentes e mais antigas.")
add_bullet(doc, "Dark Mode global com persistencia no localStorage e botao de alternancia no login e no dashboard.")
add_bullet(doc, "Exportacao rapida de resumo do prontuario para compartilhamento, com suporte a share nativo e clipboard.")

add_section_heading(doc, "3. Ferramentas usadas")
tools_table = doc.add_table(rows=1, cols=2)
tools_table.style = "Table Grid"
tools_table.autofit = False
tools_table.columns[0].width = Inches(2.1)
tools_table.columns[1].width = Inches(4.4)
hdr = tools_table.rows[0].cells
hdr[0].text = "Ferramenta"
hdr[1].text = "Uso"
for c in hdr:
    set_cell_shading(c, "D9EEF3")
    set_cell_margins(c)
    for p in c.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            style_run(r, bold=True, size=10)

tools = [
    ("React", "Interface e gerenciamento de estados."),
    ("React Router", "Roteamento e rotas dinamicas."),
    ("Axios", "Consumo da API do JSON Server."),
    ("Tailwind CSS", "Estilizacao da interface."),
    ("React Toastify", "Feedback visual de sucesso e erro."),
    ("JSON Server", "Simulacao do backend da aplicacao."),
    ("IA / Codex", "Apoio na analise, implementacao e revisao do codigo."),
]
for tool, use in tools:
    row = tools_table.add_row().cells
    row[0].text = tool
    row[1].text = use
    for idx, cell in enumerate(row):
        set_cell_margins(cell)
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                style_run(r, size=9.5)
            if idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_section_heading(doc, "4. Principais pontos do codigo")
add_body_paragraph(
    doc,
    "PatientDetails concentra o fluxo principal: busca paciente, consultas e exames pela rota /paciente/:id, trata loading e erro, permite editar e excluir registros e atualiza a interface imediatamente sem recarregar a pagina.",
)
add_body_paragraph(
    doc,
    "PatientSearchFilters foi criado para evitar repeticao da logica de busca avancada e manter as listagens consistentes em diferentes telas.",
)
add_body_paragraph(
    doc,
    "ThemeContext controla o modo claro/escuro e grava a preferencia do usuario. O arquivo index.css aplica as variaveis visuais globais do tema.",
)
add_body_paragraph(
    doc,
    "A exportacao rapida do resumo do prontuario usa navigator.share quando disponivel e faz fallback para clipboard, garantindo compatibilidade maior entre navegadores.",
)

add_section_heading(doc, "5. Prints da aplicacao")
for img_name, caption in [
    ("01-dashboard-dark-toggle.png", "Figura 1. Dashboard com modo escuro ativado."),
    ("02-prontuarios-busca-avancada.png", "Figura 2. Busca avancada na listagem de prontuarios."),
    ("03-paciente-detalhes.png", "Figura 3. Tela de detalhes do paciente com historico, ordenacao e exportacao."),
]:
    img_path = ASSETS_DIR / img_name
    if img_path.exists():
        add_image(doc, img_path)
        add_caption(doc, caption)

add_section_heading(doc, "6. Observacoes finais")
add_body_paragraph(
    doc,
    "O projeto foi validado com npm run lint e npm run build. As melhorias foram implementadas sem remover funcionalidades anteriores, preservando o comportamento existente do sistema.",
)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("Relatorio de melhorias - Clinica | Pagina ")
add_page_number(footer)

doc.save(OUTPUT_DOCX)
print(str(OUTPUT_DOCX))
